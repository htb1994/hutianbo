from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/yangcong_city_manager_self_distillation_skill.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_paragraph(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run("\n" + body)
    doc.add_paragraph()


def apply_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    apply_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "自我蒸馏 Skill 作业"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    add_page_number(section.footer.paragraphs[0])

    doc.add_paragraph("自我蒸馏 Skill 作业", style="Title")
    sub = doc.add_paragraph("主题：城市经理代理商入校获客与 C 端转化复盘", style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_callout(
        doc,
        "一句话概括",
        "这份 skill 蒸馏的是我作为洋葱学园城市经理管理代理商时的真实渠道业务流程：通过入校获客获取学生与家长线索，再围绕家长触达、体验/测评和 C 端成交做复盘、诊断和动作推进。",
    )

    doc.add_heading("一、蒸馏对象来自我的真实工作流程", level=1)
    add_paragraph(
        doc,
        "我这次自我蒸馏的对象，来自我作为洋葱学园城市经理负责渠道业务的真实工作流程。我的日常工作主要是对接区域代理商，推动代理商通过入校场景完成获客，再把学校场景中获取到的学生和家长线索转化为 C 端用户。",
    )
    add_paragraph(
        doc,
        "这个流程不是简单地“管理代理商”或“催业绩”，而是要持续判断代理商在“入校获客—线索沉淀—家长触达—体验/测评转化—C 端成交”这条链路中卡在哪个环节。",
    )
    add_paragraph(
        doc,
        "不同代理商的问题并不一样。有的代理商入校资源强，但后端 C 端转化弱；有的能进校，但活动组织粗糙，线索沉淀质量低；有的线索量够，但家长跟进不及时；有的销售团队不能把洋葱学园的产品价值讲清楚。因此，我把自己的工作流程蒸馏成一个可复用的 skill，用于代理商周复盘、问题诊断、沟通准备和下周动作制定。",
    )

    doc.add_heading("二、自我蒸馏 Skill 草稿", level=1)
    add_paragraph(doc, "Skill 名称：school-entry-agent-conversion-review")
    add_paragraph(
        doc,
        "触发描述：当城市经理需要复盘教育渠道代理商的入校获客、家长线索跟进和 C 端转化表现时使用。",
    )
    doc.add_heading("适用场景", level=2)
    add_paragraph(
        doc,
        "用于城市经理复盘代理商通过入校获客并完成 C 端转化的业务表现。适用于代理商周复盘、月度复盘、拜访准备、问题诊断、会议纪要整理和下周行动计划制定。",
    )
    doc.add_heading("触发方式", level=2)
    add_bullets(
        doc,
        [
            "需要复盘代理商本周入校获客效果。",
            "需要分析代理商 C 端转化不达标的原因。",
            "需要准备与代理商老板或团队的沟通。",
            "需要判断代理商是资源问题、执行问题，还是转化能力问题。",
            "需要向上级汇报城市渠道业务进展。",
            "需要制定下周代理商陪跑动作。",
        ],
    )
    doc.add_heading("工作流程", level=2)
    add_numbered(
        doc,
        [
            "收集事实数据：本周入校次数、覆盖学校数量、触达学生数量、获取家长线索数量、有效家长沟通数量、体验/测评预约数量、实际体验/测评数量、C 端成交数量、回款金额和上周承诺动作完成情况。",
            "判断业务卡点：区分入校不足、入校质量低、线索沉淀差、家长触达弱、体验转化弱、C 端成交弱和代理商管理问题。",
            "判断代理商类型：分为 A 类、B 类、C 类和风险类，明确下一步是加资源、陪跑、谨慎投入还是预警。",
            "输出沟通策略：先确认已完成动作，再指出最关键的业务卡点，不同时提太多问题，每次只抓一个主矛盾。",
            "形成下周行动计划：明确负责人、数量目标、时间节点和复盘方式。",
        ],
    )

    doc.add_heading("三、我的做法与判断标准", level=1)
    add_paragraph(
        doc,
        "我的做法不是只看最终成交，而是拆开完整业务链路看：入校次数 → 学生触达 → 家长线索 → 家长沟通 → 体验/测评 → C 端成交。",
    )
    add_paragraph(
        doc,
        "如果成交差，我不会直接判断“代理商不努力”，而是先判断问题在哪一段。入校次数少，说明前端资源或校方关系有问题；入校次数够但线索少，说明活动设计和留资动作有问题；线索够但体验少，说明家长触达和邀约弱；体验够但成交少，说明产品价值表达和成交跟进弱。",
    )
    add_callout(
        doc,
        "什么算做对了",
        "复盘之后，代理商知道下周具体要做什么，而不是只听到“加强转化”“提高执行”这种空泛要求。好的复盘必须能落到数量、动作、负责人和检查时间，并且下次复盘能够验证本次动作是否有效。",
    )

    doc.add_heading("四、真实或模拟运行证据", level=1)
    doc.add_heading("1. 初稿", level=2)
    add_paragraph(
        doc,
        "我最初的复盘写法是：代理商A本周成交不达标，下周需要加强入校获客，提高家长转化，争取完成目标。",
    )
    add_paragraph(
        doc,
        "这个版本的问题是太空泛。它没有说明代理商到底卡在入校、线索沉淀、家长触达，还是 C 端成交，也没有形成可检查的下周动作。",
    )

    doc.add_heading("2. 第一次运行", level=2)
    add_paragraph(doc, "我用这个 skill 对代理商A做了一次模拟复盘。")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [Inches(2.3), Inches(4.2)])
    headers = ["指标", "本周数据"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
    data = [
        ("本周入校次数", "3次"),
        ("覆盖学校", "2所"),
        ("触达学生", "480人"),
        ("获取家长线索", "96条"),
        ("有效家长沟通", "52条"),
        ("体验/测评预约", "16组"),
        ("实际体验/测评", "11组"),
        ("C 端成交", "3单"),
        ("成交金额", "15000元"),
        ("上周承诺动作", "完成3次入校活动，实际完成3次"),
    ]
    for key, value in data:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()
    add_paragraph(
        doc,
        "第一次分析结果：代理商A本周入校动作完成度较好，说明前端入校资源和执行意愿是存在的。但从480名学生触达到96条家长线索，线索沉淀率约20%；从96条线索到11组实际体验，体验转化偏弱；最终成交3单，说明主要卡点不在“能不能入校”，而在入校后的线索沉淀、家长触达和 C 端转化承接。",
    )

    doc.add_heading("3. 纠偏", level=2)
    add_paragraph(
        doc,
        "第一次运行后，我发现初稿仍然不够完整：它分析了数据，但没有判断代理商类型，也没有说明城市经理应该怎么介入。所以我更新了 skill，加入代理商分类。",
    )
    class_table = doc.add_table(rows=1, cols=2)
    class_table.style = "Table Grid"
    set_table_width(class_table, [Inches(1.4), Inches(5.1)])
    for i, text in enumerate(["类型", "判断标准"]):
        cell = class_table.cell(0, i)
        set_cell_shading(cell, "F2F4F7")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
    for key, value in [
        ("A类", "入校资源强，执行强，转化好，适合加资源和复制经验。"),
        ("B类", "有入校资源，但转化弱，需要陪跑话术、流程和数据管理。"),
        ("C类", "入校资源弱，执行也弱，需要谨慎投入。"),
        ("风险类", "连续无法入校，或入校后没有有效沉淀，需要预警。"),
    ]:
        row = class_table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()
    add_paragraph(
        doc,
        "根据这个标准，代理商A更适合归为 B 类：有入校资源，也愿意执行，但后端转化流程和团队能力偏弱。",
    )

    doc.add_heading("4. 更新后复验", level=2)
    add_paragraph(
        doc,
        "更新后的输出：代理商A当前不是入校资源不足，而是入校后的线索沉淀、家长触达和 C 端转化承接偏弱。建议将其归为 B 类代理商，下周重点进行转化链路陪跑。",
    )
    add_bullets(
        doc,
        [
            "完成不少于3次入校活动。",
            "获取家长线索不少于120条。",
            "所有家长线索在24小时内完成首轮触达。",
            "体验/测评预约不少于20组，实际体验/测评不少于15组。",
            "C 端成交目标不少于5单。",
        ],
    )
    add_paragraph(
        doc,
        "城市经理介入动作：陪代理商复盘一次入校活动后的线索整理流程；检查家长触达话术，重点优化“为什么孩子需要洋葱学园”的表达；跟进高意向家长转化节奏，明确体验后24小时内二次跟进；周五复盘入校触达数、线索沉淀率、体验预约率、实际体验率和成交率。",
    )
    add_paragraph(
        doc,
        "复验结果：更新后的版本比初稿更接近真实业务。它不是泛泛要求代理商提高业绩，而是明确判断代理商类型、定位业务卡点，并形成了下周可执行、可检查、可复盘的动作。",
    )

    doc.add_heading("五、哪些能交给 Codex，哪些必须人工判断", level=1)
    split = doc.add_table(rows=1, cols=2)
    split.style = "Table Grid"
    set_table_width(split, [Inches(3.25), Inches(3.25)])
    for i, text in enumerate(["可以交给 Codex", "必须人工判断"]):
        cell = split.cell(0, i)
        set_cell_shading(cell, "F2F4F7")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
    row = split.add_row()
    row.cells[0].text = "整理代理商周报和沟通纪要；汇总入校、触达、线索、体验、成交数据；根据漏斗数据初步识别卡点；生成沟通提纲、下周行动计划和城市渠道业务汇报。"
    row.cells[1].text = "代理商老板是否重视业务；代理商团队是否有执行能力；学校关系是否稳定；入校活动质量是否真实有效；代理商提出的问题是真困难还是借口；是否继续投入资源、加大扶持或进行预警；关键沟通中的语气、关系维护和谈判节奏。"
    for cell in row.cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.add_heading("六、总结", level=1)
    add_paragraph(
        doc,
        "这份 skill 不是一个通用销售模板，而是来自我作为洋葱学园城市经理的真实渠道业务流程。它沉淀的是我管理代理商时的判断方法：不只看代理商有没有入校，也不只看最终成交，而是完整拆解“入校获客—线索沉淀—家长触达—体验/测评转化—C 端成交”的业务链路。",
    )
    add_paragraph(
        doc,
        "通过这次自我蒸馏，我把原本依赖经验的复盘动作，变成了一套可以持续复用的个人工作流。初稿只是笼统评价代理商业绩，第一次运行后发现缺少代理商分类和城市经理介入动作，于是我补充了 A/B/C/风险类判断，并在复验中形成了更具体的下周计划。",
    )
    add_paragraph(
        doc,
        "Codex 可以帮助我整理材料、分析数据和生成复盘文本，但真正的业务判断、代理商关系判断、资源投入判断和关键谈判，仍然必须由我自己完成。这也是这个 skill 的边界和价值：让 AI 帮我提高复盘效率，但不替代我作为城市经理的业务判断。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
